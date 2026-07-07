#!/usr/bin/env python3
"""
将 CLDN18.2 临床试验数据生成 MD / DOCX / PDF / XLSX / HTML 五种格式报告。
从 enriched JSON 数据读取，包含完整字段：招募状态、最近更新、PI、联络方式、中国可报名医院及联系人。
输出采用中英文对照格式，方便国内外用户阅读。
"""

import json
import os
import re
from datetime import datetime

BASE_DIR = os.path.dirname(__file__)
OUTPUT_PARENT = os.path.join(BASE_DIR, "..", "outputs")
os.makedirs(OUTPUT_PARENT, exist_ok=True)

ENRICHED_JSON = os.path.join(OUTPUT_PARENT, "cldn18_2_enriched.json")

report_date = datetime.now().strftime("%Y-%m-%d")

# ── 搜索条件（用于子目录和文件命名） ──
SEARCH_KEYWORD = "CLDN18.2"
SEARCH_STATUS = "RECRUITING"
# 子目录名
SUB_DIR_NAME = f"{SEARCH_KEYWORD}_{SEARCH_STATUS}"
OUTPUT_DIR = os.path.join(OUTPUT_PARENT, SUB_DIR_NAME)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# 文件基础名
FILE_BASENAME = f"{SEARCH_KEYWORD}_{SEARCH_STATUS}"

# ── 中英文翻译映射表 ──────────────────────────────────────

# 临床阶段翻译
PHASE_CN = {
    "PHASE1": "I期",
    "PHASE2": "II期",
    "PHASE3": "III期",
    "PHASE4": "IV期",
    "EARLY_PHASE1": "早期I期",
    "NOT_APPLICABLE": "不适用",
}

def trans_phase(phase_en):
    """翻译临床阶段，如 'PHASE1/PHASE2' → 'I/II期 (PHASE1/PHASE2)'"""
    if not phase_en or phase_en == "Not specified":
        return "未指定 (Not specified)"
    parts = phase_en.split("/")
    cn_parts = [PHASE_CN.get(p, p) for p in parts]
    return f"{'/'.join(cn_parts)} ({phase_en})"

# 常见癌种/适应症翻译
CONDITION_TRANSLATIONS = {
    "Gastric Adenocarcinoma": "胃腺癌",
    "Gastroesophageal Junction Adenocarcinoma": "胃食管结合部腺癌",
    "Gastric or Gastroesophageal Junction Adenocarcinoma": "胃或胃食管结合部腺癌",
    "Pancreatic Ductal Adenocarcinoma": "胰腺导管腺癌",
    "Pancreatic Cancer": "胰腺癌",
    "Biliary Tract Cancer": "胆道癌",
    "Solid Tumor": "实体瘤",
    "Advanced Solid Tumor": "晚期实体瘤",
    "Malignant Solid Tumor": "恶性实体瘤",
    "Advanced Malignant Solid Tumor": "晚期恶性实体瘤",
    "Colorectal Cancer": "结直肠癌",
    "Esophageal Adenocarcinoma": "食管腺癌",
    "Esophageal Cancer": "食管癌",
    "Breast Cancer": "乳腺癌",
    "Non-small Cell Lung Cancer": "非小细胞肺癌",
    "Ovarian Cancer": "卵巢癌",
    "Peritoneal Metastasis": "腹膜转移",
    "Brain Metastasis": "脑转移",
    "Metastatic": "转移性",
    "Locally Advanced": "局部晚期",
    "Unresectable": "不可切除",
    "Resectable": "可切除",
    "Advanced": "晚期",
    "First-line": "一线",
    "Second-line": "二线",
    "Neoadjuvant": "新辅助",
    "Adjuvant": "辅助",
    "Perioperative": "围手术期",
    "HER2-positive": "HER2阳性",
    "HER2-negative": "HER2阴性",
    "CLDN18.2-positive": "CLDN18.2阳性",
    "PD-L1": "PD-L1",
    "MSI-H": "MSI-H（高度微卫星不稳定）",
    "dMMR": "dMMR（错配修复缺陷）",
}

def trans_conditions(conditions_en):
    """翻译适应症描述，保持中英对照"""
    if not conditions_en:
        return "—"
    # 尝试整句匹配
    if conditions_en in CONDITION_TRANSLATIONS:
        cn = CONDITION_TRANSLATIONS[conditions_en]
        return f"{cn} ({conditions_en})"
    # 按分隔符拆分逐段翻译
    parts = re.split(r' \| ', conditions_en)
    cn_parts = []
    for part in parts:
        part = part.strip()
        if part in CONDITION_TRANSLATIONS:
            cn_parts.append(CONDITION_TRANSLATIONS[part])
        else:
            cn_parts.append(part)
    cn_text = " | ".join(cn_parts)
    if cn_text != conditions_en:
        return f"{cn_text}（{conditions_en}）"
    return conditions_en

# 已知 CLDN18.2 靶向药物的中文名（后续可扩展）
DRUG_CN = {
    "Spevatamig (PT886)": "Spevatamig (PT886，CLDN18.2 ADC)",
    "LM-302": "LM-302（CLDN18.2 ADC）",
    "SHR-A1904": "SHR-A1904（CLDN18.2 ADC）",
    "IBI343": "IBI343（CLDN18.2 ADC）",
    "SG1906": "SG1906（CLDN18.2 ADC）",
    "RC118": "RC118（CLDN18.2 ADC）",
    "SKB315": "SKB315（CLDN18.2 ADC）",
    "JS107": "JS107（CLDN18.2 ADC）",
    "Zolbetuximab": "Zolbetuximab（CLDN18.2 单抗）",
    "AZD0901": "AZD0901（CLDN18.2 ADC）",
    "AZD5863": "AZD5863（CLDN18.2 ADC）",
    "AZD4360": "AZD4360（CLDN18.2 ADC）",
    "BDC-4182": "BDC-4182（CLDN18.2 ADC）",
    "ASP2138": "ASP2138（CLDN18.2/CD3 双抗）",
    "ASP546C": "ASP546C（CLDN18.2 双抗）",
    "CT041": "CT041（CLDN18.2 CAR-T）",
    "TST001": "TST001（CLDN18.2 单抗）",
    "Givastomig": "Givastomig（CLDN18.2 单抗）",
    "DA-3501": "DA-3501（CLDN18.2 双抗）",
    "Sonesitatug vedotin": "Sonesitatug vedotin（CLDN18.2 ADC）",
    "QLS31905": "QLS31905（CLDN18.2 双抗）",
    "XNW27011": "XNW27011（CLDN18.2 ADC）",
    "SIBP-A18": "SIBP-A18（CLDN18.2 ADC）",
    "XKDCT225": "XKDCT225（CLDN18.2 CAR-T）",
}

def annotate_drug(title):
    """在标题中标注药物中文名"""
    for en, cn in DRUG_CN.items():
        if en.lower() in title.lower():
            return title.replace(en, cn)
    return title

# ── LLM 翻译集成 ──────────────────────────────────────────
# 尝试使用 translator.py 翻译标题和适应症（如有 LLM 配置）
# 无 LLM 时静默降级为静态映射

def _try_llm_translate(trials):
    """尝试用 LLM 翻译标题和适应症，失败则返回空字典。"""
    try:
        from translator import translate_fields
        from config_loader import load_config
    except ImportError:
        return {}

    # 检查是否有 LLM 配置，无则快速跳过
    config = load_config()
    if not config.get("llm", {}).get("api_key"):
        return {}

    print("\n🌐 正在尝试 LLM 翻译标题和适应症（中英对照）...")
    print("   （如无需翻译可 Ctrl+C 跳过，自动降级为静态映射）")
    try:
        translations = translate_fields(
            items=trials,
            fields=["title", "conditions"],
            key_field="nct_id",
            domain="临床试验",
            config=config,
        )
        if translations:
            print(f"   ✅ LLM 翻译成功：{len(translations)} 项\n")
        else:
            print("   ⚠️ LLM 翻译未返回结果，使用静态映射\n")
        return translations
    except Exception as e:
        print(f"   ⚠️ LLM 翻译失败（{str(e)[:60]}），使用静态映射\n")
        return {}


# ── 加载数据 ──────────────────────────────────────────────

with open(ENRICHED_JSON, "r", encoding="utf-8") as f:
    raw_trials = json.load(f)

# LLM 翻译标题和适应症（中英对照）
llm_translations = _try_llm_translate(raw_trials)

# 将翻译合并回 trial 数据
for t in raw_trials:
    nct = t["nct_id"]
    tr = llm_translations.get(nct, {})
    # LLM 翻译的标题（如有）
    t["title_cn"] = tr.get("title", "")
    # LLM 翻译的适应症（如有）
    t["conditions_cn"] = tr.get("conditions", "")
    # 静态药物中文标注
    t["title_annotated"] = annotate_drug(t.get("title", ""))
    # 临床阶段翻译
    t["phase_cn"] = trans_phase(t.get("phase", ""))

# 按治疗类型分组（使用 NCT ID 查找 enriched 数据）
TRIAL_GROUPS = {
    "ADC药物": [
        "NCT05482893","NCT06587425","NCT06649292","NCT06350006",
        "NCT06985368","NCT05857332","NCT05205850","NCT05367635",
        "NCT07284134","NCT07584135","NCT05458219","NCT07066098",
        "NCT07483554","NCT06770439","NCT07483567","NCT07385703",
        "NCT05934331","NCT07556640","NCT07569068","NCT07450976",
        "NCT06038396","NCT06792435","NCT06519591",
    ],
    "双特异性抗体": [
        "NCT05365581","NCT07024615","NCT07481357","NCT07432295",
        "NCT07488676","NCT07431281",
    ],
    "CAR-T / CAR-NK 细胞治疗": [
        "NCT05620732","NCT06782425","NCT05911217","NCT07680257",
        "NCT07622940","NCT07480928","NCT07627711","NCT07551362",
        "NCT07066995","NCT06084286","NCT06946615","NCT07416240",
        "NCT07103668","NCT04842812","NCT03198052","NCT07523529",
    ],
    "单克隆抗体 + 联合化疗": [
        "NCT06468280","NCT06732856","NCT06962137","NCT06901531",
        "NCT07427992","NCT06902545","NCT06767449","NCT07079228",
        "NCT04495296",
    ],
    "全球 ADC（靶向 CLDN18.2）": [
        "NCT06219941","NCT05702229","NCT06921837","NCT06005493",
        "NCT06921928",
    ],
    "诊断 / 影像": [
        "NCT07301814","NCT05436093","NCT06602037","NCT07597772",
        "NCT07595237","NCT07464470",
    ],
}

# 构建 NCT_ID -> data 的快速查找表
trial_map = {t["nct_id"]: t for t in raw_trials}


def get_trials(nct_ids):
    """按 NCT ID 顺序返回试验数据，跳过缺失项。"""
    result = []
    for nid in nct_ids:
        t = trial_map.get(nid)
        if t:
            result.append(t)
    return result


def status_emoji(status):
    return {
        "RECRUITING": "🟢",
        "ACTIVE_NOT_RECRUITING": "🟡",
        "COMPLETED": "✅",
        "TERMINATED": "🔴",
        "WITHDRAWN": "⚫",
        "SUSPENDED": "🟠",
        "NOT_YET_RECRUITING": "⚪",
    }.get(status, "❓")


def status_cn(status):
    return {
        "RECRUITING": "招募中",
        "ACTIVE_NOT_RECRUITING": "暂停招募",
        "COMPLETED": "已完成",
        "TERMINATED": "已终止",
        "WITHDRAWN": "已撤回",
        "SUSPENDED": "暂停",
        "NOT_YET_RECRUITING": "尚未招募",
    }.get(status, status)


def fmt_pi(pi):
    if not pi or not pi.get("name"):
        return "未公开"
    name = pi["name"]
    aff = pi.get("affiliation", "")
    return f"{name} ({aff})" if aff else name


def fmt_contact(c):
    if not c or not c.get("name"):
        return "未公开"
    parts = [c["name"]]
    if c.get("phone"):
        parts.append(f"📞 {c['phone']}")
    if c.get("email"):
        parts.append(f"✉️ {c['email']}")
    return " | ".join(parts)


def fmt_china_hospitals(locs):
    """格式化中国可报名医院信息。"""
    if not locs:
        return "—"
    lines = []
    for loc in locs:
        parts = [loc.get("facility", ""), loc.get("city", "")]
        name_parts = []
        contact = loc.get("contact_name", "") or ""
        phone = loc.get("contact_phone", "") or ""
        email = loc.get("contact_email", "") or ""
        if contact:
            name_parts.append(contact)
        if phone:
            name_parts.append(f"📞{phone}")
        if email:
            name_parts.append(f"✉️{email}")
        contact_str = f"（{' | '.join(name_parts)}）" if name_parts else ""
        lines.append(f"  - {' '.join(parts)} {contact_str}".strip())
    return "\n".join(lines)


# ── 统计 ──
total = sum(len(ids) for ids in TRIAL_GROUPS.values())

# ═══════════════════════════════════════════════════════════
# 1. MD (Markdown)
# ═══════════════════════════════════════════════════════════

def generate_md():
    lines = []
    lines.append(f"# 🧬 CLDN18.2 靶向临床试验报告（招募中）\n")
    lines.append(f"> **报告生成日期：** {report_date}  &nbsp;|&nbsp;  **数据来源：** ClinicalTrials.gov  &nbsp;|&nbsp;  **关键词：** CLDN18.2\n")
    lines.append("")

    lines.append(f"## 📊 概要统计\n")
    lines.append(f"| 分类 | 数量 |")
    lines.append(f"|------|:----:|")
    for cat, ids in TRIAL_GROUPS.items():
        lines.append(f"| {cat} | {len(ids)} 项 |")
    lines.append(f"| **合计** | **{total} 项** |")
    lines.append("")

    for cat, ids in TRIAL_GROUPS.items():
        trials = get_trials(ids)
        if not trials:
            continue
        lines.append(f"---\n")
        lines.append(f"## 🔬 {cat}（{len(trials)} 项）\n")

        for t in trials:
            nct = t["nct_id"]
            url = t["url"]
            st = t.get("status", "")
            emoji = status_emoji(st)
            st_cn = status_cn(st)
            lu = t.get("last_update", "") or "—"
            pi = fmt_pi(t.get("pi"))
            cc = fmt_contact(t.get("central_contact"))
            china_locs = t.get("china_locations", [])
            china_hosp = fmt_china_hospitals(china_locs)

            # 中英对照：优先用 LLM 翻译，静态映射兜底
            title_cn = t.get("title_cn", "")
            title_annotated = t.get("title_annotated", "")
            phase_cn = t.get("phase_cn", "")
            cond_en = t.get("conditions", "")
            cond_cn = t.get("conditions_cn", "") or trans_conditions(cond_en)

            # 标题行：LLM 翻译（CN）+ 英文原题（EN）
            if title_cn:
                title_display = f"**CN:** {title_cn}  \n**EN:** {title_annotated}"
            else:
                title_display = f"**{title_annotated}**"

            lines.append(f"### {emoji} [{nct}]({url})\n")
            lines.append(f"{title_display}\n")
            lines.append(f"| 字段 | 内容（中文 / English） |")
            lines.append(f"|------|------------------------|")
            lines.append(f"| **招募状态** | 🟢 {st_cn} ({st}) |")
            lines.append(f"| **临床阶段** | {phase_cn} |")
            lines.append(f"| **最近更新** | {lu} |")
            lines.append(f"| **申办方** | {t.get('sponsor','') or '—'} |")
            lines.append(f"| **适应症** | {cond_cn} |")
            lines.append(f"| **PI（主要研究者）** | {pi} |")
            lines.append(f"| **中心联络** | {cc} |")
            if china_locs:
                lines.append(f"| **🏥 中国可报名医院** | |")
                for loc in china_locs:
                    fac = loc.get("facility", "")
                    city = loc.get("city", "")
                    cname = loc.get("contact_name", "")
                    cphone = loc.get("contact_phone", "")
                    contact_str = f"（{cname} | 📞{cphone}）" if cname or cphone else ""
                    lines.append(f"| | {city} {fac} {contact_str} |")
            else:
                lines.append(f"| **🏥 中国可报名医院** | — |")
            lines.append("")

    content = "\n".join(lines)
    path = os.path.join(OUTPUT_DIR, f"{FILE_BASENAME}_报告.md")
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"✅ MD 报告已生成: {path}")
    return path


# ═══════════════════════════════════════════════════════════
# 2. XLSX (Excel)
# ═══════════════════════════════════════════════════════════

def generate_xlsx():
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = openpyxl.Workbook()

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )
    wrap_align = Alignment(wrap_text=True, vertical="top")

    # ── 汇总页 ──
    ws = wb.active
    ws.title = "汇总统计"
    ws.merge_cells("A1:F1")
    ws["A1"] = "CLDN18.2 靶向临床试验报告（招募中）"
    ws["A1"].font = Font(bold=True, size=16)
    ws["A3"] = "报告生成日期"
    ws["B3"] = report_date
    ws["A4"] = "数据来源"
    ws["B4"] = "ClinicalTrials.gov"

    ws["A7"] = "分类"
    ws["B7"] = "数量"
    ws["A7"].font = Font(bold=True, color="FFFFFF")
    ws["B7"].font = Font(bold=True, color="FFFFFF")
    ws["A7"].fill = header_fill
    ws["B7"].fill = header_fill
    row = 8
    for cat, ids in TRIAL_GROUPS.items():
        ws.cell(row=row, column=1, value=cat)
        ws.cell(row=row, column=2, value=len(ids))
        row += 1
    ws.cell(row=row, column=1, value="合计").font = Font(bold=True)
    ws.cell(row=row, column=2, value=total).font = Font(bold=True)
    ws.column_dimensions["A"].width = 35
    ws.column_dimensions["B"].width = 15

    # ── 每组一个 sheet ──
    for cat, ids in TRIAL_GROUPS.items():
        safe_name = cat.replace("/", "&").replace("\\", "-")[:31]
        ws = wb.create_sheet(title=safe_name)
        headers = [
            "试验编号", "药物/方案", "适应症", "申办方",
            "招募状态", "临床阶段 (Phase)", "最近更新", "PI（主要研究者）",
            "中心联络", "中国可报名医院及联络人"
        ]
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_align
            cell.border = thin_border

        trials = get_trials(ids)
        for r, t in enumerate(trials, 2):
            nct = t["nct_id"]
            st = status_cn(t.get("status", ""))
            lu = t.get("last_update", "—")
            pi = fmt_pi(t.get("pi"))
            cc = fmt_contact(t.get("central_contact"))
            china_locs = t.get("china_locations", [])
            if china_locs:
                china_str_parts = []
                for loc in china_locs:
                    parts = []
                    parts.append(f"{loc.get('facility','')}（{loc.get('city','')}）")
                    cname = loc.get("contact_name", "")
                    cphone = loc.get("contact_phone", "")
                    if cname or cphone:
                        parts.append(f"联络人：{cname} {cphone}")
                    china_str_parts.append("；".join(parts))
                china_str = "\n".join(china_str_parts)
            else:
                china_str = "—"

            phase_cn = t.get("phase_cn", "")
            cond_cn = t.get("conditions_cn", "") or trans_conditions(t.get("conditions", ""))
            title_cn = t.get("title_cn", "")
            title_display = f"{title_cn} | {t.get('title_annotated', t.get('title',''))}" if title_cn else t.get("title_annotated", t.get("title",""))
            data = [
                nct,
                title_display,
                cond_cn,
                t.get("sponsor", ""),
                st,
                phase_cn,
                lu,
                pi,
                cc,
                china_str,
            ]
            for col, val in enumerate(data, 1):
                cell = ws.cell(row=r, column=col, value=val)
                cell.border = thin_border
                cell.alignment = wrap_align

        ws.column_dimensions["A"].width = 16
        ws.column_dimensions["B"].width = 40
        ws.column_dimensions["C"].width = 30
        ws.column_dimensions["D"].width = 22
        ws.column_dimensions["E"].width = 12
        ws.column_dimensions["F"].width = 18
        ws.column_dimensions["G"].width = 14
        ws.column_dimensions["H"].width = 30
        ws.column_dimensions["I"].width = 30
        ws.column_dimensions["J"].width = 45

    path = os.path.join(OUTPUT_DIR, f"{FILE_BASENAME}_报告.xlsx")
    wb.save(path)
    print(f"✅ XLSX 报告已生成: {path}")
    return path


# ═══════════════════════════════════════════════════════════
# 3. DOCX
# ═══════════════════════════════════════════════════════════

def generate_docx():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    title = doc.add_heading("🧬 CLDN18.2 靶向临床试验报告（招募中）", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run(f"报告生成日期：{report_date}").font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.add_run("  |  数据来源：ClinicalTrials.gov  |  关键词：CLDN18.2")

    doc.add_paragraph()

    doc.add_heading("📊 概要统计", level=1)
    # Add table manually
    table = doc.add_table(rows=len(TRIAL_GROUPS) + 2, cols=2)
    table.style = "Light Shading Accent 1"
    table.cell(0, 0).text = "分类"
    table.cell(0, 1).text = "数量"
    for i, (cat, ids) in enumerate(TRIAL_GROUPS.items(), 1):
        table.cell(i, 0).text = cat
        table.cell(i, 1).text = f"{len(ids)} 项"
    table.cell(len(TRIAL_GROUPS) + 1, 0).text = "合计"
    table.cell(len(TRIAL_GROUPS) + 1, 1).text = f"{total} 项"
    doc.add_paragraph()

    for cat, ids in TRIAL_GROUPS.items():
        trials = get_trials(ids)
        if not trials:
            continue
        doc.add_heading(f"🔬 {cat}（{len(trials)} 项）", level=1)

        for t in trials:
            nct = t["nct_id"]
            st = t.get("status", "")
            st_cn = status_cn(st)
            emoji = status_emoji(st)
            lu = t.get("last_update", "—")
            pi = fmt_pi(t.get("pi"))
            cc = fmt_contact(t.get("central_contact"))
            china_locs = t.get("china_locations", [])

            title_cn = t.get("title_cn", "")
            title_annotated = t.get("title_annotated", "")
            phase_cn = t.get("phase_cn", "")
            cond_cn = t.get("conditions_cn", "") or trans_conditions(t.get("conditions", ""))

            if title_cn:
                doc.add_heading(f"{emoji} {nct}", level=2)
                doc.add_paragraph(f"CN: {title_cn}\nEN: {title_annotated}")
            else:
                doc.add_heading(f"{emoji} {nct} — {title_annotated}", level=2)

            # 信息表格
            info_rows = [
                ("招募状态", f"{st_cn} ({st})"),
                ("临床阶段", phase_cn),
                ("最近更新", lu),
                ("申办方", t.get("sponsor", "") or "—"),
                ("适应症", cond_cn),
                ("PI（主要研究者）", pi),
                ("中心联络", cc),
            ]
            if china_locs:
                hosp_lines = []
                for loc in china_locs:
                    line = f"  • {loc.get('city','')} {loc.get('facility','')}"
                    cname = loc.get("contact_name", "")
                    cphone = loc.get("contact_phone", "")
                    if cname or cphone:
                        line += f"（{cname} | 📞{cphone}）"
                    hosp_lines.append(line)
                info_rows.append(("中国可报名医院", "\n".join(hosp_lines) if hosp_lines else "—"))
            else:
                info_rows.append(("中国可报名医院", "—"))

            tbl = doc.add_table(rows=len(info_rows), cols=2)
            tbl.style = "Light Grid Accent 1"
            for i, (label, value) in enumerate(info_rows):
                tbl.cell(i, 0).text = label
                tbl.cell(i, 1).text = value
            doc.add_paragraph()

    path = os.path.join(OUTPUT_DIR, f"{FILE_BASENAME}_报告.docx")
    doc.save(path)
    print(f"✅ DOCX 报告已生成: {path}")
    return path


# ═══════════════════════════════════════════════════════════
# 4. PDF
# ═══════════════════════════════════════════════════════════

def generate_pdf():
    from fpdf import FPDF

    FONT_PATH = "/System/Library/AssetsV2/com_apple_MobileAsset_Font7/eb257c12d1a51c8c661b89f30eec56cacf9b8987.asset/AssetData/STHEITI.ttf"

    class PDF(FPDF):
        def header(self):
            self.set_font("zh", "", 8)
            self.cell(0, 8, "CLDN18.2 Clinical Trial Report (Recruiting)", align="C", new_x="LMARGIN", new_y="NEXT")
            self.ln(3)

        def footer(self):
            self.set_y(-15)
            self.set_font("zh", "", 7)
            self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")

    pdf = PDF(orientation="L", unit="mm", format="A4")
    pdf.add_font("zh", "", FONT_PATH, uni=True)
    pdf.add_font("zh", "B", FONT_PATH, uni=True)
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # 标题
    pdf.set_font("zh", "B", 16)
    pdf.cell(0, 15, "CLDN18.2 靶向临床试验报告（招募中）", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("zh", "", 9)
    pdf.cell(0, 8, f"报告生成日期：{report_date}", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 8, "数据来源：ClinicalTrials.gov | 关键词：CLDN18.2", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)

    # 统计
    pdf.set_font("zh", "B", 12)
    pdf.cell(0, 10, "概要统计", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("zh", "", 9)
    for cat, ids in TRIAL_GROUPS.items():
        pdf.cell(0, 7, f"  {cat}: {len(ids)} 项", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("zh", "B", 10)
    pdf.cell(0, 7, f"  总计: {total} 项", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    # 每组详情
    for cat, ids in TRIAL_GROUPS.items():
        trials = get_trials(ids)
        if not trials:
            continue

        # 检查组标题+表头+一行数据的空间，不够则翻页
        needed = 15 + len(trials) * 7 + 5
        if pdf.get_y() + needed > 185:
            pdf.add_page()

        # 分类标题
        pdf.set_font("zh", "B", 11)
        pdf.set_fill_color(68, 114, 196)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(0, 8, f"  {cat}（{len(trials)} 项）", fill=True, new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(2)

        for t in trials:
            nct = t["nct_id"]
            st = t.get("status", "")
            st_cn = status_cn(st)
            emoji = status_emoji(st)
            lu = t.get("last_update", "—")[:10]
            pi = fmt_pi(t.get("pi"))
            china_locs = t.get("china_locations", [])

            title_cn = t.get("title_cn", "")
            title_annotated = t.get("title_annotated", "")
            phase_cn = t.get("phase_cn", "")
            cond_cn = t.get("conditions_cn", "") or trans_conditions(t.get("conditions", ""))

            # 标题行（优先显示 LLM 中文翻译）
            pdf.set_font("zh", "B", 8)
            if title_cn:
                title_short = title_cn[:80] + ('...' if len(title_cn) > 80 else '')
            else:
                title_short = title_annotated[:80] + ('...' if len(title_annotated) > 80 else '')
            pdf.cell(0, 6, f"{emoji} {nct} {title_short}", new_x="LMARGIN", new_y="NEXT")

            # 详情行
            pdf.set_font("zh", "", 7)
            details = f"  状态: {st_cn}  |  阶段: {phase_cn}  |  适应症: {cond_cn[:60]}{'...' if len(cond_cn)>60 else ''}  |  更新: {lu}  |  PI: {pi}"
            pdf.cell(0, 5, details, new_x="LMARGIN", new_y="NEXT")

            # 中国医院
            if china_locs:
                pdf.set_text_color(0, 100, 0)
                for loc in china_locs:
                    fac = loc.get("facility", "")
                    city = loc.get("city", "")
                    cname = loc.get("contact_name", "")
                    cphone = loc.get("contact_phone", "")
                    loc_str = f"  🏥 {city} {fac}"
                    if cname:
                        loc_str += f" 联系人: {cname}"
                    if cphone:
                        loc_str += f" 📞{cphone}"
                    pdf.cell(0, 5, loc_str[:270], new_x="LMARGIN", new_y="NEXT")
                pdf.set_text_color(0, 0, 0)

            pdf.ln(2)

        pdf.ln(3)

    path = os.path.join(OUTPUT_DIR, f"{FILE_BASENAME}_报告.pdf")
    pdf.output(path)
    print(f"✅ PDF 报告已生成: {path}")
    return path


# ═══════════════════════════════════════════════════════════
# 5. HTML
# ═══════════════════════════════════════════════════════════

def generate_html():
    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>CLDN18.2 靶向临床试验报告</title>
<style>
  * {{ margin: 0; padding: 0; box-sizing: border-box; }}
  body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif; max-width: 1200px; margin: 0 auto; padding: 20px; background: #f5f7fa; color: #333; }}
  h1 {{ color: #1a237e; border-bottom: 3px solid #1a237e; padding-bottom: 10px; margin-bottom: 10px; }}
  h2 {{ color: #283593; margin: 24px 0 12px; background: #e8eaf6; padding: 8px 15px; border-radius: 4px; font-size: 18px; }}
  h3 {{ font-size: 15px; margin: 16px 0 8px; color: #1a237e; }}
  .meta {{ color: #666; font-size: 14px; margin-bottom: 20px; }}
  .stats {{ display: flex; flex-wrap: wrap; gap: 10px; margin: 20px 0; }}
  .stat-card {{ background: white; border-radius: 8px; padding: 12px 18px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); flex: 1 1 140px; text-align: center; }}
  .stat-card .num {{ font-size: 26px; font-weight: bold; color: #1a237e; }}
  .stat-card .label {{ font-size: 12px; color: #666; margin-top: 4px; }}
  .trial-card {{ background: white; border-radius: 8px; padding: 16px 20px; margin: 12px 0; box-shadow: 0 1px 3px rgba(0,0,0,0.1); border-left: 4px solid #283593; }}
  .trial-card.recruiting {{ border-left-color: #2e7d32; }}
  .trial-card .title {{ font-size: 14px; font-weight: bold; color: #1a237e; }}
  .trial-card .title a {{ text-decoration: none; }}
  .trial-card .title a:hover {{ text-decoration: underline; }}
  .trial-card .title-cn {{ font-size: 14px; font-weight: bold; color: #1a237e; margin-top: 2px; }}
  .info-grid {{ display: grid; grid-template-columns: auto 1fr; gap: 4px 12px; margin-top: 8px; font-size: 13px; }}
  .info-grid .label {{ color: #888; font-weight: 500; min-width: 100px; }}
  .info-grid .value {{ color: #333; }}
  .hosp-table {{ margin-top: 6px; border-collapse: collapse; font-size: 12px; width: 100%; }}
  .hosp-table th {{ background: #e8f5e9; color: #2e7d32; padding: 4px 8px; text-align: left; }}
  .hosp-table td {{ padding: 3px 8px; border-bottom: 1px solid #e0e0e0; }}
  .badge {{ display: inline-block; padding: 1px 8px; border-radius: 10px; font-size: 11px; }}
  .badge-rec {{ background: #e8f5e9; color: #2e7d32; }}
  .badge-pause {{ background: #fff8e1; color: #f57f17; }}
  .footer {{ margin-top: 40px; text-align: center; color: #999; font-size: 12px; }}
</style>
</head>
<body>
<h1>🧬 CLDN18.2 靶向临床试验报告（招募中）</h1>
<div class="meta">
  📅 {report_date} &nbsp;|&nbsp; 📡 ClinicalTrials.gov &nbsp;|&nbsp; 🔑 CLDN18.2
</div>

<div class="stats">
"""
    for cat, ids in TRIAL_GROUPS.items():
        html += f'  <div class="stat-card"><div class="num">{len(ids)}</div><div class="label">{cat}</div></div>\n'
    html += f'  <div class="stat-card" style="background:#e8eaf6;"><div class="num">{total}</div><div class="label">总计</div></div>\n'
    html += '</div>\n'

    for cat, ids in TRIAL_GROUPS.items():
        trials = get_trials(ids)
        if not trials:
            continue
        html += f'<h2>🔬 {cat}（{len(trials)} 项）</h2>\n'

        for t in trials:
            nct = t["nct_id"]
            st = t.get("status", "")
            st_cn = status_cn(st)
            emoji = status_emoji(st)
            lu = t.get("last_update", "—")[:10]
            pi = fmt_pi(t.get("pi"))
            cc = fmt_contact(t.get("central_contact"))
            china_locs = t.get("china_locations", [])

            title_cn = t.get("title_cn", "")
            title_annotated = t.get("title_annotated", "")
            phase_cn = t.get("phase_cn", "")
            cond_cn = t.get("conditions_cn", "") or trans_conditions(t.get("conditions", ""))

            badge_class = "badge-rec" if st == "RECRUITING" else "badge-pause"
            html += f'<div class="trial-card {"recruiting" if st=="RECRUITING" else ""}">'
            if title_cn:
                html += f'<div class="title">{emoji} <a href="{t["url"]}" target="_blank">{nct}</a></div>'
                html += f'<div class="title-cn">🇨🇳 {title_cn}</div>'
                html += f'<div class="title-en" style="font-size:12px;color:#666;margin-top:2px;">🇬🇧 {title_annotated}</div>'
            else:
                html += f'<div class="title">{emoji} <a href="{t["url"]}" target="_blank">{nct}</a> — {title_annotated}</div>'
            html += f'<div class="info-grid">'
            html += f'<span class="label">招募状态</span><span class="value"><span class="badge {badge_class}">{st_cn}</span></span>'
            html += f'<span class="label">临床阶段</span><span class="value">{phase_cn}</span>'
            html += f'<span class="label">最近更新</span><span class="value">{lu}</span>'
            html += f'<span class="label">申办方</span><span class="value">{t.get("sponsor","") or "—"}</span>'
            html += f'<span class="label">适应症</span><span class="value">{cond_cn}</span>'
            html += f'<span class="label">PI</span><span class="value">{pi}</span>'
            html += f'<span class="label">中心联络</span><span class="value">{cc}</span>'
            if china_locs:
                html += f'<span class="label">🏥 中国可报名医院</span><span class="value">'
                html += '<table class="hosp-table"><tr><th>城市</th><th>医院</th><th>联系人</th><th>电话</th></tr>'
                for loc in china_locs:
                    html += f'<tr><td>{loc.get("city","")}</td><td>{loc.get("facility","")}</td><td>{loc.get("contact_name","")}</td><td>{loc.get("contact_phone","")}</td></tr>'
                html += '</table></span>'
            else:
                html += f'<span class="label">🏥 中国可报名医院</span><span class="value">—</span>'
            html += '</div></div>\n'

    html += f'''<div class="footer">
<p>由 clinicaltrials-search 工具自动生成 | ClinicalTrials.gov 数据</p>
</div>
</body>
</html>'''

    path = os.path.join(OUTPUT_DIR, f"{FILE_BASENAME}_报告.html")
    with open(path, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"✅ HTML 报告已生成: {path}")
    return path


# ═══════════════════════════════════════════════════════════
# Main
# ═══════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("=" * 60)
    print("  CLDN18.2 临床试验报告生成器 v2（含详细联络信息）")
    print("=" * 60)
    print()

    generate_md()
    generate_xlsx()
    generate_docx()
    generate_pdf()
    generate_html()

    print()
    print("=" * 60)
    print(f"  全部报告已生成，存放在 {OUTPUT_DIR}")
    print("=" * 60)
